import streamlit as st
from dotenv import load_dotenv
import os
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_groq import ChatGroq
import PyPDF2
import docx
from io import BytesIO
from docx import Document
import re
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from moviepy.editor import VideoFileClip, AudioFileClip
import tempfile
import uuid
import requests
import time
from math import ceil

st.set_page_config(page_title="AI Meeting Minutes Generator", layout="centered")

def apply_custom_css():
    st.markdown("""
        <style>
        html { scroll-behavior: smooth; }
        .stApp { background-color: #1C1C1E; color: #FFFFFF; }
        .navbar {
            position: fixed; top: 0; left: 0; right: 0; height: 60px;
            background: linear-gradient(to right, #7b2cbf, #9d4edd);
            z-index: 10000; display: flex; align-items: center;
            justify-content: space-between; padding: 0 30px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.2);
        }
        .navbar .brand { color: transparent; font-size: 0px; }
        .navbar a { display: none; }
        .main-content { padding-top: 80px; }
        .gradient-header-box {
            background-color: #7b2cbf; padding: 25px;
            border-radius: 10px; text-align: center;
            margin-bottom: 30px;
        }
        .gradient-header-box h1 {
            color: white; font-size: 2.5em;
            font-weight: bold; margin: 0;
        }
        .back-to-top {
            position: fixed; bottom: 30px; right: 30px;
            background: #7b2cbf; color: white;
            padding: 10px 15px; border-radius: 25px;
            text-decoration: none; font-weight: bold;
            z-index: 9999; box-shadow: 0 4px 6px rgba(0,0,0,0.2);
        }
        .section-spacer { margin-top: 100px; }
        a[href^="#"] {
            color: #d3b6f5; text-decoration: none; font-weight: bold;
        }
        a[href^="#"]:hover {
            color: #f0d9ff; text-decoration: underline;
        }
        .info-section {
            background-color: #2A2A2E;
            border-radius: 12px;
            padding: 30px;
            margin-bottom: 30px;
        }
        .info-section h3 {
            color: #d3b6f5;
            font-size: 1.8em;
            margin-bottom: 10px;
        }
        .info-section p {
            font-size: 1.1em;
            line-height: 1.6;
        }
        </style>
    """, unsafe_allow_html=True)

apply_custom_css()

st.markdown("""<div class="navbar"><div class="brand">🧠</div><div></div></div><div class="main-content">""", unsafe_allow_html=True)

load_dotenv()
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
if not GROQ_API_KEY:
    st.error("GROQ_API_KEY not found. Please check your .env file.")
    st.stop()


def initialize_llm():
    return ChatGroq(api_key=GROQ_API_KEY, model_name="llama3-8b-8192", temperature=0.3)

def create_chain(llm, full_summary=True):
    if full_summary:
        system_msg = "You are an assistant that writes professional meeting minutes."
        human_msg = (
            "Here is the meeting transcript:\n\n{transcript}\n\n"
            "Generate professional meeting minutes with the following clearly formatted markdown sections.\n\n"
            "🗓️ Date & Time\n👥 Attendee List\n💬 Key Topics Discussed\n📝 Summary\n✅ Action Items\n📅 Deadlines and Decisions\n📌 Next Steps\n\n"
            "- Use `-` for bullet points.\n"
            "- Format action items like `Name: Task`.\n"
            "- Each section must start on a new line.\n"
            "_This summary is for reference purposes only._"
            "Do not invent attendee names. Only list real attendees explicitly mentioned in the transcript."
        )
    else:
        system_msg = "You are an assistant that summarizes parts of meeting transcripts."
        human_msg = (
            "Here is a portion of a meeting transcript:\n\n{transcript}\n\n"
            "Provide a **partial summary** of what was discussed in this chunk. "
            "Focus only on what's covered here. Do not include full meeting minutes. "
            "Bullet points or concise notes are fine."
        )

    prompt = ChatPromptTemplate.from_messages([
        ("system", system_msg),
        ("human", human_msg)
    ])
    return {"transcript": RunnablePassthrough()} | prompt | llm | StrOutputParser()


# --- Text Extraction ---
def extract_text_from_pdf(uploaded_file):
    reader = PyPDF2.PdfReader(BytesIO(uploaded_file.read()))
    return "\n".join([page.extract_text() or "" for page in reader.pages]).strip()

def extract_text_from_docx(uploaded_file):
    return "\n".join([para.text for para in docx.Document(uploaded_file).paragraphs]).strip()

# --- File Generators ---
def generate_pdf_with_reportlab(text):
    from markdown2 import markdown
    import html

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=50, leftMargin=50, topMargin=50, bottomMargin=50)
    styles = getSampleStyleSheet()
    
    # Use built-in fonts instead of custom ones
    styles.add(ParagraphStyle(name='Custom', fontName='Helvetica', fontSize=11, leading=16))
    styles.add(ParagraphStyle(name='Heading', fontName='Helvetica-Bold', fontSize=18, leading=24, spaceAfter=14, spaceBefore=14))

    elements = []
    headings = ["Date & Time", "Attendee List", "Key Topics Discussed", "Summary", "Action Items", "Deadlines and Decisions", "Next Steps"]

    for line in re.sub(r"[\U00010000-\U0010ffff]", "", text).split("\n"):
        if not line.strip():
            elements.append(Spacer(1, 12))
        elif any(h in line for h in headings):
            elements.append(Paragraph(html.escape(re.sub(r'[^\w\s&]', '', line)), styles["Heading"]))
        else:
            html_line = markdown(line).strip().replace("<p>", "").replace("</p>", "")
            elements.append(Paragraph(html_line, styles["Custom"]))

    doc.build(elements)
    buffer.seek(0)
    return buffer

def generate_docx(text):
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- Transcription ---
def transcribe_audio_with_whisper(audio_file):
    headers = {"Authorization": f"Bearer {GROQ_API_KEY}"}
    files = {"file": (audio_file.name, audio_file.read(), "application/octet-stream"), "model": (None, "whisper-large-v3")}
    for attempt in range(5):
        try:
            r = requests.post("https://api.groq.com/openai/v1/audio/transcriptions", headers=headers, files=files)
            if r.status_code == 200:
                return r.json().get("text", "")
            elif r.status_code == 429:
                msg = r.json().get("error", {}).get("message", "")
                if "try again in" in msg:
                    wait_sec = float(msg.split("try again in")[-1].split("s")[0].strip()) + 1
                    st.warning(f"Rate limit hit. Waiting {wait_sec:.1f} seconds before retrying...")
                    time.sleep(wait_sec)
                    continue
                st.error(f"Rate limit: {msg}")
                return ""
            else:
                st.error(f"Transcription failed: {r.status_code}: {r.text}")
                return ""
        except Exception as e:
            st.warning(f"Retrying after error: {e}")
            time.sleep(2)
    return ""

# --- Chunked Audio Transcription ---
def split_audio_and_transcribe(audio_path, chunk_duration_sec=300):
    duration = get_audio_duration(audio_path)
    if not duration:
        st.error("Could not determine audio duration.")
        return ""

    transcript_parts = []
    total_chunks = ceil(duration / chunk_duration_sec)

    for i in range(total_chunks):
        start = i * chunk_duration_sec
        end = min((i + 1) * chunk_duration_sec, duration)
        chunk_path = os.path.join(tempfile.gettempdir(), f"chunk_{uuid.uuid4()}.mp3")
        try:
            with AudioFileClip(audio_path).subclip(start, end) as clip:
                clip.write_audiofile(chunk_path, codec='libmp3lame', verbose=False, logger=None)
            with open(chunk_path, "rb") as f:
                part = transcribe_audio_with_whisper(f)
                transcript_parts.append(part.strip())
        except Exception as e:
            st.warning(f"A chunk failed: {e}")
        finally:
            if os.path.exists(chunk_path):
                os.remove(chunk_path)
        time.sleep(1)

    return "\n".join(transcript_parts).strip()


def get_audio_duration(audio_path):
    try:
        with AudioFileClip(audio_path) as clip:
            return clip.duration
    except Exception as e:
        st.error(f"Could not read audio duration: {e}")
        return 0

def count_tokens(text):
    return len(text.split())

def split_transcript_by_tokens(transcript, max_tokens=3000):
    words = transcript.split()
    chunks = []
    for i in range(0, len(words), max_tokens):
        chunk = " ".join(words[i:i + max_tokens])
        chunks.append(chunk)
    return chunks

def extract_audio_from_uploaded_video(file):
    video_path = os.path.join(tempfile.gettempdir(), f"{uuid.uuid4()}.mp4")
    audio_path = video_path.replace(".mp4", ".mp3")
    with open(video_path, "wb") as f:
        f.write(file.read())
    try:
        with VideoFileClip(video_path) as video:
            video.audio.write_audiofile(audio_path)
        return audio_path
    except Exception as e:
        st.error(f"Video processing failed: {e}")
        return None

# --- Main App ---
def main():
    st.markdown("""<div class="gradient-header-box" id="ai-home"><h1>AI Meeting Minutes Generator</h1></div>""", unsafe_allow_html=True)

    with st.sidebar:
        st.header("⚙️ Settings & Info")
        st.markdown("📌 Paste a transcript or upload a file to begin.")
        st.markdown("🔒 Your data is processed locally and securely.")
        st.markdown("---")
        st.markdown("🔗 Quick Links")
        st.markdown('<a href="#ai-about">📖 About</a>', unsafe_allow_html=True)
        st.markdown('<a href="#ai-contact">📬 Contact</a>', unsafe_allow_html=True)

    transcript = ""
    uploaded_audio_path = None

    with st.expander("📄 Step 1: Provide Meeting Transcript"):
        method = st.radio("Choose input method", ["Paste Transcript", "Upload File", "Upload Audio/Video"])

        if method == "Paste Transcript":
            transcript = st.text_area("Paste your meeting transcript below", height=300)

        elif method == "Upload File":
            file = st.file_uploader("Upload a PDF or DOCX file", type=["pdf", "docx"])
            if file:
                if file.name.endswith(".pdf"):
                    transcript = extract_text_from_pdf(file)
                else:
                    transcript = extract_text_from_docx(file)

        elif method == "Upload Audio/Video":
            file = st.file_uploader("Upload MP3 or MP4", type=["mp3", "mp4"])
            if file:
                if file.type == "video/mp4":
                    uploaded_audio_path = extract_audio_from_uploaded_video(file)
                else:
                    uploaded_audio_path = os.path.join(tempfile.gettempdir(), f"{uuid.uuid4()}.mp3")
                    with open(uploaded_audio_path, "wb") as f:
                        f.write(file.read())

    if st.button("🚀 Generate Meeting Minutes"):
        if not transcript.strip() and not uploaded_audio_path:
            st.warning("Please provide or upload a transcript or audio/video file.")
            return

        if uploaded_audio_path and not transcript.strip():
            st.info("🔍 Transcribing uploaded audio/video...")
            audio_duration = get_audio_duration(uploaded_audio_path)
            chunk_duration = 300  # 5 minutes
            total_chunks = ceil(audio_duration / chunk_duration)
            progress = st.progress(0, text="Starting transcription...")

            transcript_parts = []
            for i in range(total_chunks):
                start = i * chunk_duration
                end = min((i + 1) * chunk_duration, audio_duration)
                chunk_path = os.path.join(tempfile.gettempdir(), f"chunk_{uuid.uuid4()}.mp3")
                try:
                    with AudioFileClip(uploaded_audio_path).subclip(start, end) as clip:
                        clip.write_audiofile(chunk_path, codec='libmp3lame', verbose=False, logger=None)
                    with open(chunk_path, "rb") as f:
                        part = transcribe_audio_with_whisper(f)
                        transcript_parts.append(part.strip())
                except Exception as e:
                    st.warning(f"A chunk failed: {e}")
                finally:
                    if os.path.exists(chunk_path):
                        os.remove(chunk_path)
                progress.progress((i + 1) / total_chunks, f"Processing chunk {i + 1}/{total_chunks}...")

            os.remove(uploaded_audio_path)
            transcript = "\n".join(transcript_parts).strip()


        if not transcript.strip():
            st.error("Transcript could not be prepared. Please try again.")
            return

        with st.spinner("Generating professional minutes..."):
            llm = initialize_llm()
            token_limit = 3000
            token_count = count_tokens(transcript)

            if token_count <= token_limit:
                chain = create_chain(llm, full_summary=True)
                result = chain.invoke(transcript)
            else:
                st.info("Processing a lengthy transcript. This may take a moment...⏳")
                chunks = split_transcript_by_tokens(transcript, max_tokens=token_limit)
                partial_results = []
                chunk_chain = create_chain(llm, full_summary=False)

                with st.spinner("🧠 Processing transcript chunks..."):
                    for chunk in chunks:
                        partial_result = chunk_chain.invoke(chunk)
                        partial_results.append(partial_result)
                        time.sleep(1.2)

                combined_summary = "\n".join(partial_results)

                # Final full summary from combined chunk summaries
                final_chain = create_chain(llm, full_summary=True)
                result = final_chain.invoke(combined_summary)


        st.success("✅ Meeting minutes generated!")
        st.subheader("📗 Generated Meeting Minutes")
        st.markdown(result)

        st.subheader("📥 Export Options")
        st.download_button("💾 Download as .txt", result, "meeting_minutes.txt", "text/plain")
        st.download_button("💾 Download as .md", result, "meeting_minutes.md", "text/markdown")
        st.download_button("💾 Download as .pdf", generate_pdf_with_reportlab(result), "meeting_minutes.pdf", "application/pdf")
        st.download_button("💾 Download as .docx", generate_docx(result), "meeting_minutes.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        

    st.markdown("""<div class="section-spacer"></div>""", unsafe_allow_html=True)
    st.markdown('<a href="#ai-home" class="back-to-top">↑ Top</a>', unsafe_allow_html=True)
    st.markdown("""
    <div class="info-section" id="ai-about">
        <h3>📖 About</h3>
        <p>This AI Meeting Minutes Generator transforms raw meeting transcripts into structured, professional minutes. It identifies attendees, key topics, summaries, action items, and decisions – saving time and increasing productivity.</p>
    </div>
    <div class="info-section" id="ai-contact">
        <h3>📬 Contact</h3>
        <p>Questions or feedback? Reach out at <a href="mailto:ai@minutes.com" style="color:#d3b6f5;">ai@minutes.com</a>.</p>
    </div>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
